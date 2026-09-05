from __future__ import annotations

import io
import mimetypes
import zipfile
from dataclasses import dataclass
from typing import Any

import fitz  # PyMuPDF
from docx import Document
from PIL import Image

from .ai import ocr_image


@dataclass
class VisualAsset:
    name: str
    data: bytes
    mime_type: str
    description: str = ""


@dataclass
class ExtractionResult:
    text: str
    images: list[VisualAsset]
    warnings: list[str]


def _image_mime(name: str, default: str = "image/png") -> str:
    return mimetypes.guess_type(name)[0] or default


def _extract_pdf(data: bytes, name: str, client: Any, language: str, model: str, max_pages: int) -> ExtractionResult:
    doc = fitz.open(stream=data, filetype="pdf")
    text_parts: list[str] = []
    images: list[VisualAsset] = []
    warnings: list[str] = []
    seen_xrefs: set[int] = set()

    page_limit = min(len(doc), max_pages)
    if len(doc) > max_pages:
        warnings.append(f"{name}: only the first {max_pages} pages were processed.")

    for page_index in range(page_limit):
        page = doc[page_index]
        page_text = page.get_text("text").strip()
        if len(page_text) < 40 and client is not None:
            pix = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
            png = pix.tobytes("png")
            try:
                page_text = ocr_image(client, png, "image/png", language, model)
            except Exception as exc:
                warnings.append(f"{name}, page {page_index + 1}: OCR failed ({exc}).")
        text_parts.append(f"\n--- {name} / page {page_index + 1} ---\n{page_text}")

        if len(images) < 16:
            for img in page.get_images(full=True):
                xref = img[0]
                if xref in seen_xrefs:
                    continue
                seen_xrefs.add(xref)
                try:
                    info = doc.extract_image(xref)
                    raw = info.get("image")
                    ext = info.get("ext", "png")
                    if not raw:
                        continue
                    with Image.open(io.BytesIO(raw)) as pil:
                        w, h = pil.size
                    if w < 240 or h < 160:
                        continue
                    images.append(VisualAsset(f"{name}-p{page_index+1}-{xref}.{ext}", raw, f"image/{ext}"))
                    if len(images) >= 16:
                        break
                except Exception:
                    continue

    return ExtractionResult("\n".join(text_parts), images, warnings)


def _extract_docx(data: bytes, name: str) -> ExtractionResult:
    document = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in document.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    images: list[VisualAsset] = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            media = [n for n in zf.namelist() if n.startswith("word/media/")]
            for media_name in media[:16]:
                raw = zf.read(media_name)
                try:
                    with Image.open(io.BytesIO(raw)) as pil:
                        w, h = pil.size
                    if w < 240 or h < 160:
                        continue
                except Exception:
                    pass
                images.append(VisualAsset(f"{name}-{media_name.split('/')[-1]}", raw, _image_mime(media_name)))
    except zipfile.BadZipFile:
        pass

    return ExtractionResult(f"\n--- {name} ---\n" + "\n".join(parts), images, [])


def _extract_image(data: bytes, name: str, mime_type: str, client: Any, language: str, model: str) -> ExtractionResult:
    if client is None:
        return ExtractionResult(
            "",
            [VisualAsset(name, data, mime_type)],
            [f"{name}: an OpenAI API key is required to read text from image files."],
        )
    text = ocr_image(client, data, mime_type, language, model)
    return ExtractionResult(f"\n--- {name} ---\n{text}", [VisualAsset(name, data, mime_type)], [])


def extract_uploaded_files(uploaded_files: list[Any], client: Any, language: str, model: str, max_pdf_pages: int = 80) -> ExtractionResult:
    text_parts: list[str] = []
    images: list[VisualAsset] = []
    warnings: list[str] = []

    for uploaded in uploaded_files:
        data = uploaded.getvalue()
        name = uploaded.name
        lower = name.lower()
        if lower.endswith(".pdf"):
            result = _extract_pdf(data, name, client, language, model, max_pdf_pages)
        elif lower.endswith(".docx"):
            result = _extract_docx(data, name)
        elif lower.endswith((".png", ".jpg", ".jpeg", ".webp")):
            mime = getattr(uploaded, "type", None) or _image_mime(name)
            result = _extract_image(data, name, mime, client, language, model)
        else:
            warnings.append(f"Unsupported file ignored: {name}")
            continue
        text_parts.append(result.text)
        images.extend(result.images)
        warnings.extend(result.warnings)

    return ExtractionResult("\n".join(text_parts).strip(), images[:24], warnings)
